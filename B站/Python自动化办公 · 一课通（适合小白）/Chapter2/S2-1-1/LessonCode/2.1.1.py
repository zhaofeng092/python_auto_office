from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt  # 磅数
from docx.oxml.ns import qn  # 中文格式
# 以上是docx库中需要用到的部分

import time

price = input('请输入今日价格：')
company_list = ['客户1', '客户2', '客户3', '客户4', '客户5', '客户6', '客户7', '客户8', '客户9', '客户10']

today = time.strftime("%Y{y}%m{m}%d{d}", time.localtime()).format(y='年', m='月', d='日')
# 获取今日时间，整理成“年-月-日”的格式

for i in company_list:
    # 针对每个客户名生成不同文档
    document = Document()

    document.styles['Normal'].font.name = u'宋体'
    # 设置文档的基础字体
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 设置文档的基础样式


    p1 = document.add_paragraph()
    # 初始化建立第一个自然段
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 对齐方式为居中，没有这句的话默认左对齐。
    run1 = p1.add_run('关于下达%s产品价格的通知' % (today))
    # 这里是第一段的内容
    run1.font.name = '微软雅黑'
    # 设置西文字体
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    # 设置中文字体
    run1.font.size = Pt(21)
    # 设置字体大小为21磅
    run1.font.bold = True
    # 设置加粗
    p1.space_after = Pt(5)
    # 段后距离5磅
    p1.space_before = Pt(5)
    # 段前距离5磅

    p2 = document.add_paragraph()
    run2 = p2.add_run(i + '：')
    # 这里是对客户的称呼
    run2.font.name = '仿宋_GB2312'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run2.font.size = Pt(16)
    run2.font.bold = True

    p3 = document.add_paragraph()
    run3 = p3.add_run('    根据公司安排，为提供优质客户服务，我单位拟定了今日价格为%s元，特此通知。' % price)
    run3.font.name = '仿宋_GB2312'
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run3.font.size = Pt(16)
    run3.font.bold = True

    p4 = document.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run('（联系人：小杨    电话：18888888888）')
    run4.font.name = '仿宋_GB2312'
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run4.font.size = Pt(16)
    run4.font.bold = True

    document.save('%s-价格通知.docx' % i)
# 以“客户名-价格通知”作为文件名保存
