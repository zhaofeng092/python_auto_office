from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt  # 磅数
from docx.oxml.ns import qn  # 中文格式

document = Document()
document.styles['Normal'].font.name = u'黑体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

def add_context(n, context):
    locals()['p' + str(n)] = document.add_paragraph()
    locals()['p' + str(n)].alignment = WD_ALIGN_PARAGRAPH.LEFT
    locals()['run' + str(n)] = locals()['p' + str(n)].add_run(str(context))
    locals()['run' + str(n)].font.size = Pt(16)
    locals()['p' + str(n)].space_after = Pt(5)
    locals()['p' + str(n)].space_before = Pt(5)

add_context(1,'The Zen of Python, by Tim Peters')
add_context(2,'Beautiful is better than ugly.')
add_context(3,'Explicit is better than implicit.')
add_context(4,'Simple is better than complex.')
add_context(5,'Complex is better than complicated.')
add_context(6,'Flat is better than nested.')
add_context(7,'Sparse is better than dense.')
add_context(8,'Readability counts.')
add_context(9,'''Special cases aren't special enough to break the rules.''')
add_context(10,'Although practicality beats purity.')

document.save('作业.docx')
