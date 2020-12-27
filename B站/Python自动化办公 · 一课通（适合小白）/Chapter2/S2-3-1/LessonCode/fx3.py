from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt  # 磅数
from docx.oxml.ns import qn  # 中文格式
from docx.shared import Inches  # 图片尺寸
# 以上是docx库中需要用到的部分
from win32com.client import Dispatch, constants, gencache
import datetime
import os
import time
price = input('请输入今日价格：')
company_list = ['客户1', '客户2', '客户3', '客户4', '客户5', '客户6', '客户7', '客户8', '客户9', '客户10']

today = datetime.date.today().strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日')

for i in company_list:
    document = Document()

    document.styles['Normal'].font.name = u'微软雅黑'
    document.styles['Normal'].font.size = Pt(16)
    # 设置文档的基础字体
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    # 设置文档的基础样式

    document.add_picture('banner.jpg', width=Inches(6))
    # 在文件最上头插入图片作为文件红头，宽度为6英寸

    p1 = document.add_paragraph()
    # 初始化建立第一个自然段
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 对齐方式为居中，没有这句的话默认左对齐。
    run1 = p1.add_run('关于下达%s产品价格的通知' % (today))
    # 这里是第一段的内容
    run1.font.name = '微软雅黑'
    # 设置西文字体
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    # 设置中文字体
    run1.font.size = Pt(21)
    # 设置字体大小为21磅
    run1.font.bold = True
    # 设置加粗
    p1.space_after = Pt(5)
    # 段后距离5磅
    p1.space_before = Pt(5)
    # 段前距离5磅

    p2 = document.add_paragraph()
    run2 = p2.add_run(i + '：')
    # 这里是对客户的称呼
    run2.font.name = '仿宋_GB2312'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run2.font.size = Pt(16)
    run2.font.bold = True

    p3 = document.add_paragraph()
    run3 = p3.add_run('    根据公司安排，为提供优质客户服务，我单位现将价格通知如下。')
    run3.font.name = '仿宋_GB2312'
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run3.font.size = Pt(16)
    run3.font.bold = True

    table = document.add_table(rows=3, cols=3, style='Table Grid')

    table.cell(0, 0).merge(table.cell(0, 2))
    table_run1 = table.cell(0, 0).paragraphs[0].add_run('XX产品报价表')
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_run1.font.name = u'隶书'
    table_run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'隶书')

    table.cell(1, 0).text = '日期'
    table.cell(1, 1).text = '价格'
    table.cell(1, 2).text = '备注'
    table.cell(2, 0).text = today
    table.cell(2, 1).text = str(price)
    table.cell(2, 2).text = ''

    p4 = document.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run('（联系人：小杨    电话：18888888888）')
    run4.font.name = '仿宋_GB2312'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run4.font.size = Pt(16)
    run4.font.bold = True

    document.add_page_break()
    p5 = document.add_paragraph()
    run4 = p5.add_run('此处是广告')
    if os.path.exists('D:/%s-价格通知.docx' % i):
        os.remove('D:/%s-价格通知.docx' % i)
    document.save('D:/%s-价格通知.docx' % i)
    # 以“客户名-价格通知”作为文件名保存
    docx_path = 'D:/%s-价格通知.docx'% i
    pdf_path = 'D:/%s-价格通知.pdf'% i


    gencache.EnsureModule('{00020905-0000-0000-C000-000000000046}', 0, 8, 4)

    wd = Dispatch("Word.Application")

    docx = wd.Documents.Open(docx_path, ReadOnly=1)
    docx.ExportAsFixedFormat(pdf_path, constants.wdExportFormatPDF, Item=constants.wdExportDocumentWithMarkup,
                             CreateBookmarks=constants.wdExportCreateHeadingBookmarks)

    wd.Quit(constants.wdDoNotSaveChanges)
    time.sleep(5)
    os.remove('D:/%s-价格通知.docx' % i)
