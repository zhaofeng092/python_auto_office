from docx import Document
from docx.enum.style import WD_STYLE_TYPE

document = Document()
styles = document.styles

for i in styles:
    if i.type == WD_STYLE_TYPE.TABLE:
        document.add_paragraph("表格样式 :  " + i.name)
        table = document.add_table(4, 5, style=i)
        document.add_paragraph("\n\n")

document.save('所有表格样式.docx')
