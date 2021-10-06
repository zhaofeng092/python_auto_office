# -*- coding: utf-8 -*-
# @Time : 2021/03/12 9:54
# @公众号 :Python自动化办公社区
# @File : python-docx.py
# @Software: PyCharm
# @Description:
# python-docx使用文档，免费送给大家，十几页s

# 1、导入python-docx库
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('主标题', 0)
# 2、新建wrod文档、一级、二级、三级标题、自然段
p = document.add_paragraph('我是自然段 ')
# 3、设置字体格式
p.add_run('我是加粗、黑体').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('我是一级标题', level=1)
document.add_paragraph('我是自然段2', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    '我是数字序号', style='List Number'
)
# 4、在指定位置添加图片
document.add_picture('demo.jpg', width=Inches(5))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)
# 5、在指定位置添加表格
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '我是第1列'
hdr_cells[1].text = '我是第2列'
hdr_cells[2].text = '我是第3列'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('test.docx')